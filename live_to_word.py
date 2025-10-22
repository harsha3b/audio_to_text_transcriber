import queue, threading, time, datetime, pathlib, sys, numpy as np
import sounddevice as sd
from docx import Document # to write .docx
from faster_whisper import WhisperModel # to transcribe
import webrtcvad
import os, shutil

# --- Optional VAD: we'll try to import; if not available, we fall back to timing-based chunks ---
try:
    import webrtcvad
    VAD_AVAILABLE = True
except Exception:
    VAD_AVAILABLE = False

# ===== SETTINGS  =====
LANG = "en"         # "auto", "de", "en", ...
MODEL = "small"       # "tiny"/"base" faster; "medium" more accurate (slower)
DEVICE = "cpu"        # CPU is fine on your laptop
COMPUTE = "int8"      # best for CPU speed
SAMPLE_RATE = 16000
FRAME_MS = 20         # 10/20/30 ms (buffer size per audio frame).
CHUNK_SEC = 3         # how often to send a block for transcription.
DOC_DIR = pathlib.Path.home() / "Desktop" / "Journal" / "audio_to_txt"
MIC_DEVICE = None     # None = default mic; or set an integer index (see listing snippet below)

DOC_DIR.mkdir(parents=True, exist_ok=True)
model = WhisperModel(MODEL, device=DEVICE, compute_type=COMPUTE)

audio_q = queue.Queue() # holds incoming audio blocks

# 'audio_callback' pushes small int16 audio blocks into audio_q as they arrive.
def audio_callback(indata, frames, time_info, status):
    if status:
        print(status, file=sys.stderr)
    audio_q.put(indata.copy())

def read_mic():
    with sd.InputStream(samplerate=SAMPLE_RATE, channels=1, dtype="int16",
                        device=MIC_DEVICE,
                        blocksize=int(SAMPLE_RATE * FRAME_MS / 1000),
                        callback=audio_callback):
        while True:
            time.sleep(0.05)

""" If webrtcvad is available: collect_chunks_vad() checks each 20ms frame to see if it’s speech. 
It accumulates voiced frames (with a little tolerance for short silence) until ~CHUNK_SEC seconds, 
then emits one chunk. """
def collect_chunks_vad():
    vad = webrtcvad.Vad(2)  # 0..3 (3 = most aggressive)
    bytes_per_frame = int(SAMPLE_RATE * (FRAME_MS / 1000.0)) * 2  # int16 mono
    window = b""
    voiced = []
    voiced_len = 0.0
    while True:
        block = audio_q.get()
        b = block.tobytes()
        window += b
        while len(window) >= bytes_per_frame:
            frame = window[:bytes_per_frame]
            window = window[bytes_per_frame:]
            is_voiced = vad.is_speech(frame, SAMPLE_RATE)
            if is_voiced:
                voiced.append(frame)
                voiced_len += FRAME_MS/1000.0
            else:
                if voiced:
                    voiced.append(frame)
                    voiced_len += FRAME_MS/1000.0
            if voiced_len >= CHUNK_SEC:
                chunk = b"".join(voiced)
                yield np.frombuffer(chunk, dtype=np.int16)
                voiced, voiced_len = [], 0.0
""" If VAD is not available: collect_chunks_timer() 
just emits fixed-length chunks by time (simpler, a bit less clean around pauses). """
def collect_chunks_timer():
    # Simple fallback: emit fixed-size chunks regardless of speech
    bytes_per_block = int(SAMPLE_RATE * (FRAME_MS / 1000.0)) * 2
    need_bytes = int(SAMPLE_RATE * CHUNK_SEC) * 2
    buf = bytearray()
    while True:
        block = audio_q.get()
        buf.extend(block.tobytes())
        if len(buf) >= need_bytes:
            out = bytes(buf[:need_bytes])
            del buf[:need_bytes]
            yield np.frombuffer(out, dtype=np.int16)

# calls whisper to get text from the chuck of audio and concatenates the pieces
def transcribe_pcm16(pcm16: np.ndarray):
    floats = pcm16.astype(np.float32) / 32768.0
    segments, info = model.transcribe(
        floats,
        language=None if LANG == "auto" else LANG,
        vad_filter=True
    )
    return "".join(s.text for s in segments).strip()


# cretaes a doc with todays date
def doc_path_today():
    return DOC_DIR / f"{datetime.date.today().isoformat()}.docx"

# creates a  doc with todays date if it does not exist
def ensure_doc():
    p = doc_path_today()
    if not p.exists():
        doc = Document()
        doc.add_heading(f"Live Journal — {datetime.date.today().isoformat()}", 0)
        doc.save(p)
    return p



word_opened = False  # keep at top of your file (global flag)

def append_to_doc(text: str):
    if not text:
        return
    p = ensure_doc()
    tmp = p.with_suffix(".tmp.docx")

    doc = Document(p)

    # Add timestamp once at top if not already there
    if len(doc.paragraphs) == 0 or not any(h.style.name.startswith("Heading") for h in doc.paragraphs):
        now = datetime.datetime.now().strftime("%H:%M:%S")
        doc.add_heading(now, level=2)
        doc.add_paragraph("")  # spacer
        doc.add_paragraph(text)
        
    else:
        # Append text to the same last paragraph (same line)
        last_para = doc.paragraphs[-1]
        last_para.add_run(" " + text.strip())
        

    doc.save(tmp)

    try:
        shutil.move(tmp, p)
    except PermissionError:
        # print("⚠️ Word is locking the file — open the .tmp.docx copy instead.")
        print(f"{text}")
    else:
        # print(f"Appended transcript to {p}")
        print(f"{text}")

    # Open file only once for live viewing
    global word_opened
    if not word_opened:
        os.startfile(p)
        word_opened = True


def main():
    threading.Thread(target=read_mic, daemon=True).start()
    print("🎙️ Live transcription started. Speak into your mic…")
    print("    Writing to:", doc_path_today())
    collector = collect_chunks_vad if VAD_AVAILABLE else collect_chunks_timer 
    
    for chunk in collector():
        try:
            text = transcribe_pcm16(chunk)
            append_to_doc(text)
        except Exception as e:
            print("Transcription error:", e)
    

if __name__ == "__main__":
    main()
